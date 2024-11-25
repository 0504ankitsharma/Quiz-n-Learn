from docx import Document
import io

def export_to_doc(mcqs, include_answers=True):
    doc = Document()
    doc.add_heading('Multiple Choice Questions', level=1)
    
    for idx, mcq in enumerate(mcqs, 1):
        doc.add_heading(f"Question {idx} (20 points)", level=2)
        doc.add_paragraph(mcq['question'])
        for choice in mcq['choices']:
            doc.add_paragraph(choice)
        if include_answers:
            doc.add_paragraph(f"Answer: {mcq['answer']}", style='Emphasis')
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer